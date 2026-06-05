from docx import Document
from docx.shared import Pt

doc = Document("datawoordenboek.docx")


def add_heading(doc, text, level=2):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16) if level == 2 else Pt(13)
    return p


def add_simple_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    for ci, h in enumerate(headers):
        cell = table.cell(0, ci)
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, 1):
        for ci, v in enumerate(row):
            table.cell(ri, ci).text = v
    return table


# Bestand 4
doc.add_page_break()
add_heading(doc, "Bestand 4: Psychologie 2022-2023", 2)

p = doc.add_paragraph()
p.add_run("Bestandsnaam: ").bold = True
p.add_run("2022-2023 Totaalscores Psychologie.xlsx")
p = doc.add_paragraph()
p.add_run("Blad: ").bold = True
p.add_run("'Totaalscores met formules'   |   Aantal rijen: 10   |   Aantal kolommen: 9")

doc.add_paragraph(
    "Dit bestand is gebruikt voor de selectie van studenten voor de bacheloropleiding Psychologie, studiejaar 2022-2023. "
    "Het bestand is veel eenvoudiger dan het 2026-2027 bestand: alle scores zijn al geaggregeerd tot drie instrumentscores en een totaalscore. "
    "Er zijn geen individuele vakscores of keuzevakken."
)

add_heading(doc, "Instrumenten in dit bestand", 3)
add_simple_table(
    doc,
    ["Instrument", "Beschrijving", "Kolommen"],
    [
        [
            "Selectietoets",
            "Een selectietoets specifiek voor de opleiding.",
            "Toets (ruwe score), Toetsscore (percentage)",
        ],
        [
            "Matchingsvragenlijst",
            "Een vragenlijst die meet hoe goed de motivatie en verwachtingen passen bij de opleiding.",
            "Matching (ruwe score), Matchingscore (percentage)",
        ],
        [
            "Cijferlijst",
            "Een samenvatting van de cijfers op het schooldiploma.",
            "Cijferlijst (ruwe score), Cijferlijstscore (percentage)",
        ],
    ],
)

add_heading(doc, "Alle kolommen", 3)
add_simple_table(
    doc,
    ["Kolomnaam", "Type", "Beschrijving"],
    [
        ["Studentnummer", "Getal", "Uniek studentnummer van de kandidaat."],
        ["Toets", "Getal", "Ruwe score op de selectietoets."],
        ["Toetsscore", "Decimaal (%)", "Procentuele score op de selectietoets."],
        ["Matching", "Getal", "Ruwe score op de matchingsvragenlijst."],
        [
            "Matchingscore",
            "Decimaal (%)",
            "Procentuele score op de matchingsvragenlijst.",
        ],
        ["Cijferlijst", "Decimaal", "Ruwe score op de cijferlijst."],
        ["Cijferlijstscore", "Decimaal (%)", "Procentuele score op de cijferlijst."],
        ["Totaalscore", "Decimaal (%)", "Gewogen totaalscore als percentage."],
        ["Rangnummer", "Getal", "Het definitieve rangnummer van de kandidaat."],
    ],
)

# Bestand 5
doc.add_page_break()
add_heading(doc, "Bestand 5: Psychologie 2021-2022", 2)

p = doc.add_paragraph()
p.add_run("Bestandsnaam: ").bold = True
p.add_run("2021-2022 Totaalscores Psychologie.xls")
p = doc.add_paragraph()
p.add_run("Blad: ").bold = True
p.add_run("'Totaalscores met formules'   |   Aantal rijen: 10   |   Aantal kolommen: 9")
p = doc.add_paragraph()
p.add_run("Let op: ").bold = True
p.add_run("dit bestand is in het oudere .xls-formaat (Excel 97-2003).")

doc.add_paragraph(
    "Dit bestand heeft dezelfde structuur als 2022-2023, maar de kolomnamen verschillen. "
    'Waar 2022-2023 "Toets" en "Matching" gebruikt, heet het hier "TOETS" en "MATCH".'
)

add_heading(doc, "Alle kolommen", 3)
add_simple_table(
    doc,
    ["Kolomnaam", "Type", "Beschrijving"],
    [
        ["Studentnummer", "Getal", "Uniek studentnummer van de kandidaat."],
        ["TOETS", "Getal", "Ruwe score op de selectietoets."],
        ["Toets score", "Decimaal (%)", "Procentuele score op de selectietoets."],
        ["MATCH", "Getal", "Ruwe score op de matchingsvragenlijst."],
        [
            "Match score",
            "Decimaal (%)",
            "Procentuele score op de matchingsvragenlijst.",
        ],
        ["CIJF", "Decimaal", "Ruwe score op de cijferlijst."],
        ["Cijf score", "Decimaal (%)", "Procentuele score op de cijferlijst."],
        ["Totaalscore", "Decimaal (%)", "Gewogen totaalscore als percentage."],
        ["Rangnummer", "Getal", "Het definitieve rangnummer van de kandidaat."],
    ],
)

# Vergelijkingstabel
doc.add_page_break()
add_heading(doc, "Vergelijking Psychologie-bestanden over de jaren", 2)

doc.add_paragraph(
    "De drie Psychologie-bestanden hebben dezelfde drie instrumenten (selectietoets, matching, cijferlijst/schooldiploma), "
    "maar de structuur verschilt sterk. De oudere bestanden hebben al geaggregeerde scores, terwijl 2026-2027 individuele vakscores heeft."
)

add_simple_table(
    doc,
    ["Kenmerk", "2021-2022", "2022-2023", "2026-2027"],
    [
        ["Aantal kolommen", "9", "9", "46"],
        ["Bestandsformaat", ".xls", ".xlsx", ".xlsx"],
        [
            "Individuele vakscores",
            "Nee, alleen totaal",
            "Nee, alleen totaal",
            "Ja, per vak",
        ],
        ["Keuzevakken", "Niet apart", "Niet apart", "V1 t/m V10 apart"],
        [
            "Structuur scores",
            "Ruwe score + percentage",
            "Ruwe score + percentage",
            "Punten (0-5) per vak + deelscores",
        ],
    ],
)

try:
    doc.save("datawoordenboek.docx")
    print("Datawoordenboek aangevuld.")
except PermissionError:
    doc.save("datawoordenboek_v2.docx")
    print(
        "Datawoordenboek opgeslagen als datawoordenboek_v2.docx (origineel is open in Word)."
    )
